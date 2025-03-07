import threading

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse
from django.shortcuts import redirect, render
from django.views.generic import CreateView, DetailView, ListView, UpdateView
from docx import Document

from .forms import TranscribeAudioForm
from .models import TranscribeAudio
from .tasks import transcribe_audio


@login_required
def create_transcription(request):
    if request.method == "POST":
        form = TranscribeAudioForm(request.POST, request.FILES)
        if form.is_valid():
            transcription = form.save(commit=False)
            transcription.created_by = request.user
            transcription = form.save()
            transcribe_audio_thread = threading.Thread(
                target=transcribe_audio, args=(transcription.id,)
            )
            transcribe_audio_thread.start()
        return redirect("detail", slug=form.instance.slug)

    else:
        form = TranscribeAudioForm()
    return render(request, "audio/transcribe_audio.html", {"form": form})


class TranscribeAudioListView(ListView):
    model = TranscribeAudio


@login_required
def transcription_list_view(request):
    transcription_list = TranscribeAudio.objects.filter(created_by=request.user)
    return render(request, "audio/list.html", {"object_list": transcription_list})


@login_required
def transcription_detail_view(request, slug):
    transcription = TranscribeAudio.objects.get(slug=slug)
    return render(request, "audio/detail.html", {"object": transcription})


@login_required
def export_to_word(request, slug):
    # Create a new Document
    doc = Document()
    data = TranscribeAudio.objects.get(slug=slug)

    # Add a title
    doc.add_heading(f"{data.name}", 0)

    # Get the data you want to export
    document_name = data.name

    # Add data to the document
    doc.add_paragraph(f"Probability: {data.probability}")
    doc.add_paragraph(f"Duration: {data.get_duration()}")
    doc.add_paragraph(f"Description: {data.transcription}")
    doc.add_paragraph("")  # Blank line between entries

    # Prepare the response to download the Word document
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    filename = f"{document_name}.docx"
    response["Content-Disposition"] = f'attachment; filename="{filename}"'

    # Save the document to the response
    doc.save(response)
    return response
